import os
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

def add_images_to_docx(doc, image_dir, image_order, captions):
    for img, caption in zip(image_order, captions):
        img_path = os.path.join(image_dir, img)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(4))
            doc.add_paragraph(caption)
        else:
            doc.add_paragraph(f"[Image not found: {img}]")

def add_images_to_pdf(pdf, image_dir, image_order, captions):
    for img, caption in zip(image_order, captions):
        img_path = os.path.join(image_dir, img)
        if os.path.exists(img_path):
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, caption, ln=True)
            pdf.image(img_path, w=120)
        else:
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, f"[Image not found: {img}]", ln=True)

def main():
    # Document content
    doc_text = [
        "OptiFoot Software — Complete Expert Breakdown",
        "\n1. THE BIG PICTURE — What problem are we solving?\nDiabetic patients develop Peripheral Artery Disease (PAD) — blood flow to the feet decreases. This causes tissue to lose oxygen, eventually forming ulcers that can lead to amputation.\n\nThe problem? By the time a doctor sees an ulcer, it's already too late. Current methods (Doppler ultrasound, visual inspection) are expensive, subjective, and detect damage only after it's visible.\n\nOptiFoot detects the oxygen drop before the ulcer forms — using light.",
        "\n2. THE SCIENCE — How does light measure oxygen?\nBlood has two forms of haemoglobin — oxygenated (HbO₂) and deoxygenated (HHb). These two absorb light differently at different wavelengths. We exploit this.",
        "\n650 nm (red): High-power red LED — HHb absorbs a lot here, HbO₂ absorbs very little\n850 nm (near-infrared): High-power NIR LED — Both absorb similarly (reference)\n\nIf tissue has low oxygen → more HHb → it absorbs more 650nm light → the camera sees a darker reflection at 650nm compared to 850nm. We measure that ratio.\n\nThis is the same principle as a pulse oximeter on your finger — but we do it spatially across the entire foot, producing a 2D map instead of a single number.",
        "\n3. THE ALGORITHM — Beer-Lambert Law\nWe use the modified Beer-Lambert law. When light passes through tissue, absorption depends on the concentration of absorbers and their extinction coefficients.\n\nFormula implemented:\nR = ln(I650) / ln(I850)\nSpO₂ = [εHHb850 − εHbO₂850 − R⋅(εHHb650 − εHbO₂650)] / [εHHb850 − R⋅εHHb650] × 100%\n\nWhere:\n- I650, I850 = pixel intensity (how much light reflected back)\n- ε = extinction coefficient (known constants from literature)\n- R = the ratio that encodes oxygenation information\n\nThis runs per-pixel, so you get a full 2D oxygenation map of the foot — not just one number.",
        "\n4. THE SOFTWARE PIPELINE — Step by step\nStep 1: Image Capture (hardware-software interface)\nThe Raspberry Pi controls two LEDs via GPIO. It turns on the 650nm LED, waits 80ms for stable illumination, captures a frame with the NoIR camera, then switches to 850nm and captures again. This gives us two images of the same foot under different wavelengths.\n\nWhy NoIR camera? Normal cameras have an IR filter that blocks 850nm light. The NoIR (No Infrared filter) camera lets NIR light through.\n\nStep 2: Preprocessing\nThe two images may have slight misalignment because of the 80ms gap between captures. We use ECC-based image registration (Enhanced Correlation Coefficient) to align them pixel-perfectly. Then we apply Gaussian blur for noise reduction and Otsu thresholding to segment the foot from the dark background.\n\nIf they ask why alignment matters: A 1-pixel shift between the 650 and 850nm images would create false SpO₂ readings at tissue boundaries.\n\nStep 3: SpO₂ Computation\nWe apply the Beer-Lambert formula to every pixel simultaneously using NumPy vectorized operations. The output is a 2D float array where each value represents the estimated tissue oxygen saturation (0–100%) at that location.\n\nStep 4: Heatmap Visualization\nThe SpO₂ map is converted to a colour image using the JET colourmap — blue = high oxygen (healthy), red = low oxygen (at risk). We overlay contour lines around regions below critical thresholds.\n\nStep 5: Risk Scoring\nWe compute a composite risk score (0–100) from four weighted factors:\nMean SpO₂ (35%) — Overall oxygenation level\n% Critical pixels (<85%) (30%) — How much tissue is severely hypoxic\n% At-risk pixels (85-95%) (20%) — How much tissue is borderline\nLargest critical cluster (15%) — Whether low-oxygen areas are concentrated (worse) or scattered\n\nClassification: Normal (<20) → Monitor (20-40) → At Risk (40-60) → Critical (>60)",
        "\nStep 6: Storage & Temporal Tracking\nEach scan is saved to SQLite with the SpO₂ map, heatmap, and all metrics. Over multiple visits, we can show whether the patient is improving or deteriorating by comparing scan-to-scan difference maps.",
        "\n5. KEY DESIGN DECISIONS — What to highlight if asked\nWhy threshold-based scoring instead of ML?\nFor a medical device prototype, we need explainable results. A threshold-based system lets us say exactly why a score is 65 — 'because 40% of the tissue is below 85% SpO₂'. We designed the code with a strategy pattern so we can drop in a trained CNN model later without changing the rest of the pipeline.\n\nWhy PyQt5 for GUI?\nIt runs natively on Raspberry Pi, supports zoom/pan on medical images via QGraphicsView, and has a more clinical appearance than Tkinter. No internet required — everything runs on-device.\n\nWhy SQLite?\nZero-configuration, single-file database. Perfect for a point-of-care device that doesn't need a server.\n\nCan it run without the Pi?\nYes — the --demo flag activates synthetic image generation and GPIO stubs. The entire pipeline runs on any laptop for development and testing.",
        "\n6. WHAT THE DEMO IMAGES SHOW — Walk the reviewer through",
        "7. TOTAL COST JUSTIFICATION\nA hospital Doppler ultrasound costs ₹2-5 lakhs. Our entire device costs ₹11,545 — Raspberry Pi (₹6000) + NoIR camera (₹2500) + LEDs (₹1400) + PCB (₹200) + enclosure. This can be deployed in rural clinics where expensive equipment isn't available.",
        "8. ONE-LINER SUMMARY\nOptiFoot uses dual-wavelength near-infrared imaging to calculate per-pixel tissue oxygenation via the Beer-Lambert law, producing a real-time SpO₂ heatmap and risk score that detects diabetic foot complications before ulcers form — all on a ₹11,545 Raspberry Pi device."
    ]

    REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    image_dir = os.path.join(REPO_ROOT, "optifoot", "demo_output")
    docs_dir = os.path.join(REPO_ROOT, "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    image_order = [
         "03_dual_wavelength_comparison.png",
         "06_heatmap_with_colorbar.png",
         "05_heatmap_risk_zones.png",
         "07_risk_score_panel.png",
         "09_temporal_comparison_strip.png",
         "10_architecture_diagram.png"
    ]
    captions = [
         "03_dual_wavelength_comparison.png — These are the two raw captures. Left is 650nm, right is 850nm. Notice the intensity differences — that's what encodes oxygenation.",
         "06_heatmap_with_colorbar.png — After Beer-Lambert processing, here's the SpO₂ map. Blue regions have healthy oxygen levels. Red regions are hypoxic — potential ulcer formation sites.",
         "05_heatmap_risk_zones.png — We automatically delineate critical zones with contours. A clinician can see exactly where intervention is needed.",
         "07_risk_score_panel.png — The composite risk score with all metrics. This is what the clinician sees — one number plus supporting data.",
         "09_temporal_comparison_strip.png — Over multiple visits, we track changes. This shows baseline vs follow-up with a difference map.",
         "10_architecture_diagram.png — The software architecture — modular pipeline from LED control through to GUI and database."
    ]
    
    # Create Word document
    doc = Document()
    for section in doc_text:
         doc.add_paragraph(section)
    add_images_to_docx(doc, image_dir, image_order, captions)
    doc.save(os.path.join(docs_dir, "OptiFoot_Expert_Breakdown.docx"))

    # Create PDF document
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    def sanitize(text):
        return (text
            .replace("—", "-")
            .replace("–", "-")
            .replace("−", "-")
            .replace("⋅", ".")
            .replace("₹", "Rs.")
            .replace("’", "'")
            .replace("“", '"')
            .replace("”", '"')
            .replace("₂", "2")
            .replace("→", "->")
            .replace("ε", "e")
        )
    for section in doc_text:
        pdf.multi_cell(0, 10, sanitize(section))
    add_images_to_pdf(pdf, image_dir, image_order, [sanitize(c) for c in captions])
    pdf.output(os.path.join(docs_dir, "OptiFoot_Expert_Breakdown.pdf"))

if __name__ == "__main__":
    main()
